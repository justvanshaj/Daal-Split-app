import streamlit as st
from fpdf import FPDF
from PIL import Image
import io
import re
import unicodedata
import datetime
import logging
from pathlib import Path

# third-party for docx handling
try:
    import docx  # python-docx
    from docx import Document
except Exception as ex:
    docx = None
    Document = None

logger = logging.getLogger(__name__)

# ---------- CONFIG ----------
st.set_page_config(page_title="Aravally Dal Split", page_icon="favicon_split.ico")
st.title("Aravally Dal Split")

st.markdown("""
<style>
#MainMenu {visibility: hidden;}
footer {visibility: hidden;}
header {visibility: hidden;}
.result-table { width: 100%; border-collapse: collapse; margin-bottom: 10px; }
.result-table th, .result-table td { border: 1px solid #ddd; padding: 8px; text-align: left; }
.result-table th { background-color: #f2f2f2; font-weight:600; }
@media (max-width: 480px) {
  .result-table td, .result-table th { font-size: 14px; padding: 6px; }
}
</style>
""", unsafe_allow_html=True)

# ---------- DATE ----------
selected_date = st.date_input("Select Date (DD/MM/YYYY)", value=datetime.date.today())
display_date = selected_date.strftime("%d/%m/%Y")
filename_date = selected_date.strftime("%Y-%m-%d")

# ---------- INPUTS ----------
vehicle_number = st.text_input("Enter Vehicle Number:")
party_name = st.text_input("Enter Party Name:")
gaadi_type = st.radio("Select Gaadi Type:", options=["Khadi", "Poori"])

st.write("**Enter weights (grams)** — 3 decimal precision")
a = st.number_input("Daal (input)", min_value=0.000, step=0.001, format="%.3f")
b = st.number_input("Tukdi (input)", min_value=0.000, step=0.001, format="%.3f")
c = st.number_input("Red/Black (input)", min_value=0.000, step=0.001, format="%.3f")
d = st.number_input("Chhala (input)", min_value=0.000, step=0.001, format="%.3f")
e = st.number_input("Dankhal (input)", min_value=0.000, step=0.001, format="%.3f")
f = st.number_input("14 Mesh (input)", min_value=0.000, step=0.001, format="%.3f")

# ---------- IMAGE UPLOADER ----------
uploaded_file = st.file_uploader("Upload an image (jpg, jpeg, png) — this will be page 1 of the PDF", type=["jpg", "jpeg", "png"])
image_file = uploaded_file

# ---------- CALCULATIONS ----------
g = round(a + b + c + d + e + f, 3)
h = round(a * 2, 3)  # Daal grams for sheet
i = round(b * 2, 3)  # Tukdi grams for sheet
j = round(c * 2, 3)
k = round(d * 2, 3)
l = round(e * 2, 3)
m = round(f * 2, 3)
grand_total = round(h + i + j + k + l + m, 3)

# Percent (original logic preserved)
h_percent = round(h * 10, 3)
i_percent = round(i * 10, 3)
j_percent = round(j * 10, 3)
k_percent = round(k * 10, 3)
l_percent = round(l * 10, 3)
m_percent = round(m * 10, 3)

total_dal_tukdi_percent = round(h_percent + i_percent, 3)
total_4_percent = round(j_percent + k_percent + l_percent + m_percent, 3)
total_6_percent = round(total_dal_tukdi_percent + total_4_percent, 3)

# New grams totals
total_dal_tukdi_grams = round(h + i, 3)
total_4_grams = round(j + k + l + m, 3)

# ---------- RENDER RESULTS (responsive HTML table) ----------
def render_results_table():
    html = f"""
    <table class="result-table">
      <thead>
        <tr><th>Item</th><th>Grams</th><th>Percentage</th></tr>
      </thead>
      <tbody>
        <tr><td>Daal</td><td>{h} g</td><td>{h_percent} %</td></tr>
        <tr><td>Tukdi</td><td>{i} g</td><td>{i_percent} %</td></tr>
        <tr style="font-weight:700;"><td>Total (Dal + Tukdi)</td><td>{total_dal_tukdi_grams} g</td><td>{total_dal_tukdi_percent} %</td></tr>
        <tr><td>Red/Black</td><td>{j} g</td><td>{j_percent} %</td></tr>
        <tr><td>Chhala</td><td>{k} g</td><td>{k_percent} %</td></tr>
        <tr><td>Dankhal</td><td>{l} g</td><td>{l_percent} %</td></tr>
        <tr><td>14 Mesh</td><td>{m} g</td><td>{m_percent} %</td></tr>
        <tr style="font-weight:700;"><td>Total (4)</td><td>{total_4_grams} g</td><td>{total_4_percent} %</td></tr>
        <tr style="font-weight:700;"><td>Grand Total for Sheet</td><td>{grand_total} g</td><td>{total_6_percent} %</td></tr>
      </tbody>
    </table>
    """
    st.markdown(html, unsafe_allow_html=True)

st.subheader("Grand Total")
st.write(f"Grand Total (sum of raw inputs): **{g} g**")
st.subheader("Details — grams and percentage (rows keep alignment on mobile)")
render_results_table()

# ---------- TEMPLATE PATHS ----------
TEMPLATE_DOCX_PATH = Path("report_template.docx")  # <-- your .docx template
FILLED_DOCX_PATH = Path(f"filled_report_{filename_date}.docx")
FILLED_PDF_PATH = Path(f"filled_report_{filename_date}.pdf")

# ---------- DOCX PLACEHOLDER REPLACEMENT ----------
def replace_placeholders_in_docx(template_path: Path, output_path: Path, data_map: dict):
    """
    Load a .docx, replace placeholders ({{KEY}} or {KEY}) with values from data_map,
    and save output docx to output_path.
    """
    if Document is None:
        raise RuntimeError("python-docx (docx) library is required. Install with 'pip install python-docx'")

    doc = Document(str(template_path))

    # Create combined mapping of supported token styles
    # We'll replace {{KEY}} and {KEY} both
    replace_map = {}
    for k, v in data_map.items():
        replace_map[f"{{{{{k}}}}}"] = str(v)  # {{KEY}}
        replace_map[f"{{{k}}}"] = str(v)      # {KEY}

    # Replace in paragraphs
    for p in doc.paragraphs:
        text = p.text
        new_text = text
        for token, val in replace_map.items():
            if token in new_text:
                new_text = new_text.replace(token, val)
        if new_text != text:
            # simple replacement: replace whole paragraph text (runs/styles may be lost)
            p.clear()  # remove existing runs
            p.add_run(new_text)

    # Replace in tables (cells)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text
                new_text = cell_text
                for token, val in replace_map.items():
                    if token in new_text:
                        new_text = new_text.replace(token, val)
                if new_text != cell_text:
                    # setting cell.text will overwrite paragraphs in cell
                    cell.text = new_text

    # headers & footers (optional)
    try:
        for section in doc.sections:
            header = section.header
            for p in header.paragraphs:
                text = p.text
                new_text = text
                for token, val in replace_map.items():
                    if token in new_text:
                        new_text = new_text.replace(token, val)
                if new_text != text:
                    p.clear()
                    p.add_run(new_text)
    except Exception:
        # not critical
        pass

    doc.save(str(output_path))
    return output_path

# ---------- helper to extract text from docx (for fallback PDF rendering) ----------
def extract_text_from_docx(path: Path):
    if Document is None:
        return ""
    doc = Document(str(path))
    lines = []
    for p in doc.paragraphs:
        stripped = p.text.strip()
        if stripped:
            lines.append(stripped)
    # also include table text rows
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(cells))
    return "\n".join(lines)

# ---------- attempt to convert docx -> pdf using docx2pdf or libreoffice if available ----------
def try_convert_docx_to_pdf(input_docx: Path, output_pdf: Path):
    # Try docx2pdf first (Windows/Mac)
    try:
        import docx2pdf
        try:
            docx2pdf.convert(str(input_docx), str(output_pdf))
            return True, "docx2pdf"
        except Exception as e:
            logger.debug("docx2pdf failed: %s", e)
    except Exception:
        pass

    # Try LibreOffice (soffice) headless conversion
    try:
        import subprocess, shutil
        soffice = shutil.which("soffice") or shutil.which("libreoffice")
        if soffice:
            # run soffice to convert
            cmd = [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(output_pdf.parent), str(input_docx)]
            subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            # produced file will be same name with .pdf in outdir
            expected = output_pdf.parent / (input_docx.stem + ".pdf")
            if expected.exists():
                expected.rename(output_pdf)
                return True, "libreoffice"
    except Exception as e:
        logger.debug("libreoffice conversion failed: %s", e)

    return False, None

# ---------- PDF generation using FPDF ----------
def pdf_from_image_and_text(image_file, report_text, output_bytes=True):
    pdf = FPDF()

    # First page = image (landscape)
    pdf.add_page(orientation='L')
    if image_file:
        try:
            img = Image.open(image_file)
            buf = io.BytesIO()
            img.convert("RGB").save(buf, format="JPEG")
            buf.seek(0)
            pdf.image(buf, x=10, y=10, w=270)
        except Exception as e:
            pdf.set_font("Arial", "B", 16)
            pdf.cell(0, 10, f"Error placing image: {e}", ln=True, align='C')
    else:
        pdf.set_font("Arial", "B", 16)
        pdf.cell(0, 10, "No Image Provided", ln=True, align='C')

    # Second page = report_text (already filled)
    pdf.add_page(orientation='P')
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)
    # write report text preserving newlines
    for line in str(report_text).splitlines():
        pdf.multi_cell(0, 8, line)
    if output_bytes:
        return pdf.output(dest='S').encode('latin-1')
    else:
        return pdf

# ---------- DATA MAP / PLACEHOLDERS ----------
data_map = {
    "DATE": display_date,
    "VEHICLE": vehicle_number,
    "PARTY": party_name,
    "GAADI": gaadi_type,
    "DAAL": h,
    "TUKDI": i,
    "TOTAL_DTG": total_dal_tukdi_grams,   # your grams placeholder
    "TOTAL_DTP": total_dal_tukdi_percent, # your percent placeholder
    "REDBLACK": j,
    "CHHALA": k,
    "DANKHAL": l,
    "MES14": m,
    "TOTAL_4_GRAMS": total_4_grams,
    "TOTAL_4": total_4_percent,
    "GRAND_TOTAL": grand_total,
    "TOTAL_6": total_6_percent,
    "DAAL_PERC": h_percent,
    "TUKDI_PERC": i_percent,
    "REDBLACK_PERC": j_percent,
    "CHHALA_PERC": k_percent,
    "DANKHAL_PERC": l_percent,
    "MES14_PERC": m_percent
}

# ---------- MAIN: Fill template, produce docx and PDF ----------
if st.button("Generate PDF"):
    # 1) If docx template exists, create filled docx
    if TEMPLATE_DOCX_PATH.exists() and Document is not None:
        try:
            filled_docx = replace_placeholders_in_docx(TEMPLATE_DOCX_PATH, FILLED_DOCX_PATH, data_map)
            st.success(f"Filled DOCX saved: {FILLED_DOCX_PATH.name}")
            with open(FILLED_DOCX_PATH, "rb") as fh:
                st.download_button("📥 Download filled DOCX", data=fh, file_name=FILLED_DOCX_PATH.name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        except Exception as e:
            st.error(f"Failed to fill DOCX template: {e}")
            filled_docx = None
    else:
        if not TEMPLATE_DOCX_PATH.exists():
            st.warning(f"Docx template not found at {TEMPLATE_DOCX_PATH}. Falling back to plain-text report in PDF.")
        if Document is None:
            st.warning("python-docx library not available; install 'python-docx' to use a .docx template.")
        filled_docx = None

    # 2) Try to convert filled docx to PDF (preferred to preserve styling)
    created_pdf_bytes = None
    if filled_docx:
        success, method = try_convert_docx_to_pdf(filled_docx, FILLED_PDF_PATH)
        if success:
            st.success(f"Converted filled DOCX to PDF using: {method}")
            # return the produced PDF file bytes
            with open(FILLED_PDF_PATH, "rb") as fh:
                pdf_bytes = fh.read()
            created_pdf_bytes = pdf_bytes
            st.download_button("📥 Download converted PDF", data=pdf_bytes, file_name=FILLED_PDF_PATH.name, mime="application/pdf")
        else:
            st.info("Could not auto-convert DOCX -> PDF (docx2pdf or libreoffice not available). Will render the filled docx content into the PDF (content kept; original styling lost).")

    # 3) If conversion not available, extract text and render into PDF
    if created_pdf_bytes is None:
        if filled_docx and Document is not None:
            report_text = extract_text_from_docx(filled_docx)
        else:
            # create a simple text-based report (same fallback as earlier)
            lines = [
                "Dal Split Report",
                f"Date: {data_map['DATE']}",
                f"Vehicle Number: {data_map['VEHICLE']}",
                f"Party Name: {data_map['PARTY']}",
                f"Gaadi Type: {data_map['GAADI']}",
                "",
                "Details (grams):",
                f"Daal: {data_map['DAAL']} gm",
                f"Tukdi: {data_map['TUKDI']} gm",
                f"Total (Dal + Tukdi): {data_map['TOTAL_DTG']} gm",
                f"Red/Black: {data_map['REDBLACK']} gm",
                f"Chhala: {data_map['CHHALA']} gm",
                f"Dankhal: {data_map['DANKHAL']} gm",
                f"14 Mesh: {data_map['MES14']} gm",
                f"Total (4): {data_map['TOTAL_4_GRAMS']} gm",
                f"Grand Total for Sheet: {data_map['GRAND_TOTAL']} gm",
                "",
                "Details (percentage):",
                f"Daal: {data_map['DAAL_PERC']} %",
                f"Tukdi: {data_map['TUKDI_PERC']} %",
                f"Total (Dal + Tukdi): {data_map['TOTAL_DTP']} %",
                f"Red/Black: {data_map['REDBLACK_PERC']} %",
                f"Chhala: {data_map['CHHALA_PERC']} %",
                f"Dankhal: {data_map['DANKHAL_PERC']} %",
                f"14 Mesh: {data_map['MES14_PERC']} %",
                f"Total (4): {data_map['TOTAL_4']} %",
                f"Total (6): {data_map['TOTAL_6']} %"
            ]
            report_text = "\n".join(lines)

        # create PDF bytes using FPDF with the uploaded image as page 1
        try:
            pdf_bytes = pdf_from_image_and_text(image_file, report_text, output_bytes=True)
            st.success("PDF generated (text-rendered from the filled template).")
            fname = f"{filename_date}_{unicodedata.normalize('NFKD', vehicle_number).encode('ascii','ignore').decode('ascii').strip().replace(' ','_')}_{unicodedata.normalize('NFKD',party_name).encode('ascii','ignore').decode('ascii').strip().replace(' ','_')}_{gaadi_type}_gaadi.pdf"
            st.download_button("📥 Download PDF", data=pdf_bytes, file_name=fname, mime="application/pdf")
        except Exception as e:
            st.error(f"Failed to generate PDF: {e}")
